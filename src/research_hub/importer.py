"""Local file import pipeline for non-DOI documents."""

from __future__ import annotations

import hashlib
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml

from research_hub.clusters import ClusterRegistry, slugify
from research_hub.dedup import DedupHit, DedupIndex
from research_hub.security import atomic_write_text, safe_join, validate_slug

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: dict[str, str] = {
    "pdf": "pdf",
    "md": "markdown",
    "markdown": "markdown",
    "txt": "txt",
    "docx": "docx",
    "url": "url",
}

_HASH_KEY_PREFIX = "hash:"
_BODY_PREVIEW_LIMIT = 5000
_SUMMARY_LIMIT = 500


@dataclass
class ImportEntry:
    path: Path
    slug: str = ""
    status: str = ""
    error: str = ""
    note_path: Path | None = None
    source_kind: str = ""


@dataclass
class ImportReport:
    folder: Path
    cluster_slug: str
    entries: list[ImportEntry] = field(default_factory=list)
    dry_run: bool = False

    @property
    def imported_count(self) -> int:
        return sum(1 for entry in self.entries if entry.status == "imported")

    @property
    def skipped_count(self) -> int:
        return sum(1 for entry in self.entries if entry.status.startswith("skipped"))

    @property
    def failed_count(self) -> int:
        return sum(1 for entry in self.entries if entry.status == "failed")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "PDF extraction requires pdfplumber. Install: pip install 'research-hub-pipeline[import]'"
        ) from exc

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n\n".join(parts).strip()


def _extract_markdown(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end > 0:
            text = text[end + 5 :]
    return text.strip()


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "DOCX extraction requires python-docx. Install: pip install 'research-hub-pipeline[import]'"
        ) from exc

    doc = docx.Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()


def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except OSError as exc:
        raise RuntimeError(f"could not read {path}: {exc}") from exc


def _extract_url(path: Path) -> str:
    try:
        import requests
        from readability import Document as ReadabilityDocument
    except ImportError as exc:  # pragma: no cover - dependency-gated
        raise RuntimeError(
            "URL extraction requires requests + readability-lxml. Install: pip install 'research-hub-pipeline[import]'"
        ) from exc

    first_line = next((line.strip() for line in path.read_text(encoding="utf-8").splitlines() if line.strip()), "")
    if not first_line.startswith(("http://", "https://")):
        raise ValueError(f"{path}: first non-empty line must be a URL")
    response = requests.get(
        first_line,
        timeout=30,
        headers={"User-Agent": "research-hub/0.31"},
    )
    response.raise_for_status()
    return ReadabilityDocument(response.text).summary()


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "markdown": _extract_markdown,
    "docx": _extract_docx,
    "txt": _extract_txt,
    "url": _extract_url,
}


def _derive_title(text: str, fallback: Path) -> str:
    for line in text.splitlines()[:50]:
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    return fallback.stem.replace("_", " ").replace("-", " ").strip() or fallback.stem


def _filename_slug(path: Path) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", path.stem.lower()).strip("-")
    return base[:64] or "imported-file"


def _derive_slug(title: str, source_path: Path) -> str:
    candidate = slugify(title)
    candidate = re.sub(r"[^a-z0-9_-]+", "-", candidate.lower()).strip("-_")
    return (candidate[:64] if candidate else _filename_slug(source_path)) or "imported-file"


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()[:16]


def _truncate_body(text: str) -> str:
    if len(text) <= _BODY_PREVIEW_LIMIT:
        return text
    return text[:_BODY_PREVIEW_LIMIT] + "\n\n*(truncated; see raw_path for the source file)*"


def _document_frontmatter(
    *,
    slug_value: str,
    title: str,
    source_kind: str,
    cluster_slug: str,
    raw_path: Path,
    summary: str,
) -> dict[str, Any]:
    try:
        from research_hub.document import Document

        doc = Document(
            slug=slug_value,
            title=title,
            source_kind=source_kind,
            topic_cluster=cluster_slug,
            ingestion_source="import-folder",
            raw_path=str(raw_path),
            summary=summary,
        )
        if hasattr(doc, "to_frontmatter"):
            frontmatter = dict(doc.to_frontmatter())
        else:
            frontmatter = {
                "title": title,
                "slug": slug_value,
                "source_kind": source_kind,
                "topic_cluster": cluster_slug,
                "ingestion_source": "import-folder",
                "raw_path": str(raw_path),
                "summary": summary,
            }
    except Exception:
        frontmatter = {
            "title": title,
            "slug": slug_value,
            "source_kind": source_kind,
            "topic_cluster": cluster_slug,
            "ingested_at": "",
            "ingestion_source": "import-folder",
            "labels": [],
            "tags": [],
            "raw_path": str(raw_path),
            "summary": summary,
        }

    frontmatter.setdefault("slug", slug_value)
    frontmatter["title"] = title
    frontmatter["source_kind"] = source_kind
    frontmatter["topic_cluster"] = cluster_slug
    frontmatter["ingestion_source"] = "import-folder"
    frontmatter["raw_path"] = str(raw_path)
    frontmatter["summary"] = summary
    frontmatter.setdefault("labels", [])
    frontmatter.setdefault("tags", [])
    frontmatter.setdefault("cluster_queries", [])
    frontmatter.setdefault("status", "unread")
    return frontmatter


def _render_document_md(
    *,
    slug_value: str,
    title: str,
    source_kind: str,
    cluster_slug: str,
    raw_path: Path,
    text: str,
) -> str:
    frontmatter = _document_frontmatter(
        slug_value=slug_value,
        title=title,
        source_kind=source_kind,
        cluster_slug=cluster_slug,
        raw_path=raw_path,
        summary=text[:_SUMMARY_LIMIT],
    )
    yaml_str = yaml.safe_dump(frontmatter, allow_unicode=True, sort_keys=False)
    body = _truncate_body(text).strip()
    return f"---\n{yaml_str}---\n\n{body}\n"


def _maybe_create_cluster(cfg: Any, cluster_slug: str, *, dry_run: bool) -> None:
    registry = ClusterRegistry(cfg.clusters_file)
    if registry.get(cluster_slug) is not None or dry_run:
        return
    registry.create(
        query=cluster_slug,
        name=cluster_slug.replace("-", " ").title(),
        slug=cluster_slug,
    )
    logger.info("auto-created cluster: %s", cluster_slug)


def _load_dedup(path: Path) -> DedupIndex:
    return DedupIndex.load(path) if path.exists() else DedupIndex()


def _hash_key(content_hash: str) -> str:
    return f"{_HASH_KEY_PREFIX}{content_hash}"


def import_folder(
    cfg: Any,
    folder: Path | str,
    *,
    cluster_slug: str,
    extensions: tuple[str, ...] | None = None,
    skip_existing: bool = True,
    use_graphify: bool = False,
    dry_run: bool = False,
) -> ImportReport:
    cluster_slug = validate_slug(cluster_slug, field="cluster_slug")
    folder_path = Path(folder).expanduser().resolve()
    if not folder_path.is_dir():
        raise ValueError(f"folder not found: {folder_path}")

    requested_extensions = extensions or ("pdf", "md", "txt", "docx", "url")
    normalized_extensions = tuple(ext.lower().lstrip(".") for ext in requested_extensions)

    _maybe_create_cluster(cfg, cluster_slug, dry_run=dry_run)
    dedup_path = cfg.research_hub_dir / "dedup_index.json"
    dedup = _load_dedup(dedup_path)
    cluster_raw_dir = safe_join(cfg.raw, cluster_slug)
    if not dry_run:
        cluster_raw_dir.mkdir(parents=True, exist_ok=True)

    report = ImportReport(folder=folder_path, cluster_slug=cluster_slug, dry_run=dry_run)

    for path in sorted(folder_path.rglob("*")):
        if not path.is_file():
            continue

        ext = path.suffix.lower().lstrip(".")
        if ext not in normalized_extensions or ext not in SUPPORTED_EXTENSIONS:
            report.entries.append(ImportEntry(path=path, status="skipped_unsupported"))
            continue

        source_kind = SUPPORTED_EXTENSIONS[ext]
        entry = ImportEntry(path=path, source_kind=source_kind)
        try:
            text = _EXTRACTORS[source_kind](path)
        except Exception as exc:
            entry.status = "failed"
            entry.error = f"extraction failed: {exc}"
            report.entries.append(entry)
            continue

        title = _derive_title(text, path)
        entry.slug = _derive_slug(title, path)
        content_hash = _content_hash(text)
        if skip_existing and dedup.title_to_hits.get(_hash_key(content_hash)):
            entry.status = "skipped_duplicate"
            report.entries.append(entry)
            continue

        markdown = _render_document_md(
            slug_value=entry.slug,
            title=title,
            source_kind=source_kind,
            cluster_slug=cluster_slug,
            raw_path=path,
            text=text,
        )

        if not dry_run:
            note_path = cluster_raw_dir / f"{entry.slug}.md"
            if note_path.exists():
                note_path = cluster_raw_dir / f"{entry.slug}-{content_hash[:8]}.md"
            atomic_write_text(note_path, markdown)
            entry.note_path = note_path
            dedup.title_to_hits[_hash_key(content_hash)] = [
                DedupHit(
                    source="importer",
                    title=title,
                    obsidian_path=str(note_path),
                )
            ]

        entry.status = "imported"
        report.entries.append(entry)

    if not dry_run and report.imported_count > 0:
        dedup.save(dedup_path)

    if use_graphify and not dry_run:
        _run_graphify(report, folder_path)

    return report


def _run_graphify(report: ImportReport, folder_path: Path) -> None:
    try:
        from research_hub.graphify_bridge import (
            map_to_subtopics,
            parse_graphify_communities,
            run_graphify,
        )
    except ImportError:
        logger.warning("graphify_bridge not available; skipping --use-graphify step")
        return

    try:
        graph_json = run_graphify(folder_path)
        communities = parse_graphify_communities(graph_json)
        assignments = map_to_subtopics(
            communities,
            [entry.path for entry in report.entries if entry.status == "imported"],
        )
        for entry in report.entries:
            if entry.status != "imported" or entry.note_path is None:
                continue
            subtopics = assignments.get(str(entry.path), [])
            if subtopics:
                _add_subtopics_frontmatter(entry.note_path, subtopics)
    except Exception as exc:  # pragma: no cover - best effort bridge
        logger.warning("graphify integration failed: %s", exc)


def _add_subtopics_frontmatter(note_path: Path, subtopics: list[str]) -> None:
    text = note_path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        return
    end = text.find("\n---\n", 4)
    if end < 0:
        return
    frontmatter = yaml.safe_load(text[4:end]) or {}
    if not isinstance(frontmatter, dict):
        return
    frontmatter["subtopics"] = list(subtopics)
    yaml_str = yaml.safe_dump(frontmatter, allow_unicode=True, sort_keys=False)
    body = text[end + 5 :]
    note_path.write_text(f"---\n{yaml_str}---\n{body}", encoding="utf-8")


__all__ = ["import_folder", "ImportEntry", "ImportReport", "SUPPORTED_EXTENSIONS"]
